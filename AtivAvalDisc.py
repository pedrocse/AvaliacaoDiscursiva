# -*- coding: utf-8 -*-
from __future__ import annotations

import os
import re
import smtplib
import time
import uuid
from datetime import datetime, timedelta
from email.message import EmailMessage
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Optional, Tuple
from zipfile import ZIP_DEFLATED, ZipFile

import pandas as pd
import psycopg2
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from psycopg2.extras import Json, RealDictCursor


# ============================================================
# CONFIG
# ============================================================

APP_TITLE = "Avaliação Discursiva - Taxonomia de Bloom"

BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
DOCX_DIR = DATA_DIR / "docx"
IMAGE_DIR = DATA_DIR / "question_images"

TAXONOMIA_OPTIONS = ["Baixa", "Média", "Alta"]
AVALIACAO_OPTIONS = ["A1", "A2", "A3(A1)", "A3(A2)"]
BIMESTRE_OPTIONS = ["1B", "2B"]

STATUS_OPTIONS = [
    "RECEBIDA",
    "EM REVISAO",
    "APROVADA",
    "CORRECAO SOLICITADA",
    "ENVIADA AO ALUNO",
]

ANSWER_STATUS_OPTIONS = [
    "ENVIADA",
    "CORRIGIDA",
]

DIMENSOES_AVALIACAO = [
    "Elaborou a resposta com letra legível, coerência e coesão textuais; teve capricho.",
    "Obedeceu a taxonomia de Bloom.",
    "Desenvolveu a resposta conforme a sequência de verbos solicitada no enunciado da questão.",
    "Atendeu ao conteúdo de conhecimento cognitivo previsto no gabarito.",
    "Utilizou adequadamente o espaço destinado para a resposta.",
]

INSTRUCOES_PROVA = [
    "Leia atentamente os textos contidos neste caderno e responda as questões propostas.",
    "Coloque seu nome em todas as folhas.",
    "Utilize apenas caneta azul ou preta.",
    "A prova terá duração máxima de 2 (duas) horas.",
    "A interpretação dos enunciados faz parte da aferição de conhecimentos e da avaliação.",
    "Serão considerados na avaliação: capricho, letra legível, organização, clareza textual, atendimento aos comandos na sequência em que foram solicitados.",
    "As questões devem ser respondidas utilizando-se somente os espaços para elas determinado.",
    "É proibido o empréstimo de material escolar e o uso de corretivos na avaliação.",
    "É vedado o uso de equipamentos eletrônicos, celulares e outros meios de comunicação durante a prova, exceto quando o professor da disciplina autorizar seu uso.",
]

MAX_LOGIN_ATTEMPTS = 5
LOCK_MINUTES = 10


# ============================================================
# SECRETS / ENV
# ============================================================

def get_secret(key: str, default: Any = None) -> Any:
    try:
        if key in st.secrets:
            return st.secrets[key]
        return default
    except Exception:
        return default


SMTP_HOST = os.getenv("SMTP_HOST") or get_secret("SMTP_HOST", "")
SMTP_PORT = int(os.getenv("SMTP_PORT") or get_secret("SMTP_PORT", 587))
SMTP_USER = os.getenv("SMTP_USER") or get_secret("SMTP_USER", "")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD") or get_secret("SMTP_PASSWORD", "")
SMTP_FROM = os.getenv("SMTP_FROM") or get_secret("SMTP_FROM", SMTP_USER)

# Neste app local, o arquivo .streamlit/secrets.toml tem prioridade.
# Assim, variáveis antigas do Windows usadas para Supabase não sobrescrevem a configuração local.
PGHOST = get_secret("PGHOST", None) or os.getenv("PGHOST") or "localhost"
PGPORT = int(get_secret("PGPORT", None) or os.getenv("PGPORT") or 5432)
PGDATABASE = get_secret("PGDATABASE", None) or os.getenv("PGDATABASE") or "avaliacao_discursiva"
PGUSER = get_secret("PGUSER", None) or os.getenv("PGUSER") or "postgres"
PGPASSWORD = get_secret("PGPASSWORD", None) or os.getenv("PGPASSWORD") or ""

ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD") or get_secret("ADMIN_PASSWORD", "")


# ============================================================
# HELPERS
# ============================================================

def ensure_dirs() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)


def count_chars(s: str) -> int:
    return len(s or "")


def is_email_valid(email: str) -> bool:
    email = (email or "").strip()
    return bool(re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email))


def safe_slug(s: str, max_len: int = 40) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"[^a-z0-9_.-]+", "_", s)
    return s[:max_len].strip("_")


def safe_filename(name: str, max_len: int = 120) -> str:
    name = (name or "").strip()
    name = re.sub(r"[^\w\-. ]+", "_", name, flags=re.UNICODE)
    name = re.sub(r"\s+", " ", name).strip()
    return name[:max_len] if max_len else name


def allowed_image_extension(filename: str) -> bool:
    return Path(filename or "").suffix.lower() in [".png", ".jpg", ".jpeg"]


def has_supported_image_signature(file_bytes: bytes) -> bool:
    return file_bytes.startswith(b"\x89PNG\r\n\x1a\n") or file_bytes.startswith(b"\xff\xd8")


def save_question_image(uploaded_file, evaluation_id: str, question_idx: int) -> tuple[str, str]:
    file_bytes = uploaded_file.getbuffer().tobytes()
    if not has_supported_image_signature(file_bytes):
        raise ValueError("Imagem inválida. Use um arquivo PNG ou JPG/JPEG real.")

    original_name = safe_filename(uploaded_file.name or f"questao_{question_idx}.png")
    suffix = Path(original_name).suffix.lower() or ".png"
    image_name = f"{evaluation_id[:8]}_q{question_idx}_{uuid.uuid4().hex[:8]}{suffix}"
    image_path = IMAGE_DIR / image_name
    image_path.write_bytes(file_bytes)
    return image_name, str(image_path)


def read_file_bytes(path: str | Path) -> bytes:
    with open(path, "rb") as f:
        return f.read()


def now_str() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def generate_access_code() -> str:
    return uuid.uuid4().hex[:8].upper()


# ============================================================
# ADMIN AUTH
# ============================================================

def _now() -> datetime:
    return datetime.now()


def admin_auth_guard(admin_password: str) -> bool:
    if "admin_authed" not in st.session_state:
        st.session_state.admin_authed = False
    if "admin_attempts" not in st.session_state:
        st.session_state.admin_attempts = 0
    if "admin_lock_until" not in st.session_state:
        st.session_state.admin_lock_until = None

    if not admin_password:
        st.warning("ADMIN_PASSWORD não está configurada. O painel admin está sem proteção.")
        return True

    if st.session_state.admin_authed:
        return True

    lock_until = st.session_state.admin_lock_until
    if lock_until is not None and _now() < lock_until:
        remaining = lock_until - _now()
        mins = int(remaining.total_seconds() // 60)
        secs = int(remaining.total_seconds() % 60)
        st.error(f"Acesso ao admin bloqueado temporariamente. Tente novamente em {mins:02d}:{secs:02d}.")
        return False

    st.info("Acesso restrito. Digite a senha do admin.")
    pwd = st.text_input("Senha do admin", type="password", key="admin_pwd")

    if st.button("Entrar", use_container_width=True):
        if pwd == admin_password:
            st.session_state.admin_authed = True
            st.session_state.admin_attempts = 0
            st.session_state.admin_lock_until = None
            st.success("Acesso liberado.")
            st.rerun()
        else:
            st.session_state.admin_attempts += 1
            remaining_attempts = MAX_LOGIN_ATTEMPTS - st.session_state.admin_attempts
            if remaining_attempts <= 0:
                st.session_state.admin_lock_until = _now() + timedelta(minutes=LOCK_MINUTES)
                st.session_state.admin_attempts = 0
                st.error(f"Senha incorreta. Acesso bloqueado por {LOCK_MINUTES} minutos.")
            else:
                st.error(f"Senha incorreta. Tentativas restantes: {remaining_attempts}")

    return False


# ============================================================
# POSTGRESQL
# ============================================================

def db_is_configured() -> bool:
    return bool(PGHOST and PGPORT and PGDATABASE and PGUSER)


def db_connect():
    ensure_dirs()
    last_error = None

    for _ in range(3):
        try:
            return psycopg2.connect(
                host=PGHOST,
                port=PGPORT,
                dbname=PGDATABASE,
                user=PGUSER,
                password=PGPASSWORD,
                cursor_factory=RealDictCursor,
                connect_timeout=10,
            )
        except psycopg2.OperationalError as e:
            last_error = e
            time.sleep(1)

    raise last_error


def db_init() -> None:
    if not db_is_configured():
        st.error("PostgreSQL não está configurado. Confira PGHOST, PGPORT, PGDATABASE e PGUSER.")
        st.stop()

    ddl = """
    CREATE TABLE IF NOT EXISTS evaluations (
        id TEXT PRIMARY KEY,
        created_at TIMESTAMP NOT NULL,
        updated_at TIMESTAMP NOT NULL,
        nome_docente TEXT NOT NULL,
        email_docente TEXT NOT NULL,
        email_aluno TEXT NOT NULL DEFAULT '',
        area_conhecimento TEXT NOT NULL,
        avaliacao TEXT NOT NULL,
        bimestre TEXT NOT NULL,
        turma TEXT NOT NULL,
        periodo_letivo TEXT NOT NULL,
        data_avaliacao TEXT NOT NULL,
        valor_total TEXT NOT NULL,
        titulo TEXT NOT NULL,
        access_code TEXT UNIQUE,
        questoes JSONB NOT NULL,
        docx_filename TEXT NOT NULL,
        docx_path TEXT NOT NULL,
        status TEXT NOT NULL DEFAULT 'RECEBIDA',
        observacao_admin TEXT NOT NULL DEFAULT ''
    );
    """

    student_answers_ddl = """
    CREATE TABLE IF NOT EXISTS student_answers (
        id TEXT PRIMARY KEY,
        evaluation_id TEXT NOT NULL REFERENCES evaluations(id) ON DELETE CASCADE,
        access_code TEXT NOT NULL,
        created_at TIMESTAMP NOT NULL,
        updated_at TIMESTAMP NOT NULL,
        nome_aluno TEXT NOT NULL,
        email_aluno TEXT NOT NULL,
        respostas JSONB NOT NULL,
        docx_filename TEXT NOT NULL,
        docx_path TEXT NOT NULL,
        status TEXT NOT NULL DEFAULT 'ENVIADA',
        observacao_admin TEXT NOT NULL DEFAULT ''
    );
    """

    indexes = [
        "CREATE INDEX IF NOT EXISTS idx_evaluations_status ON evaluations(status);",
        "CREATE INDEX IF NOT EXISTS idx_evaluations_created_at ON evaluations(created_at DESC);",
        "CREATE INDEX IF NOT EXISTS idx_evaluations_email_docente ON evaluations(email_docente);",
        "CREATE INDEX IF NOT EXISTS idx_evaluations_email_aluno ON evaluations(email_aluno);",
        "CREATE INDEX IF NOT EXISTS idx_evaluations_area ON evaluations(area_conhecimento);",
        "CREATE UNIQUE INDEX IF NOT EXISTS idx_evaluations_access_code ON evaluations(access_code);",
        "CREATE INDEX IF NOT EXISTS idx_student_answers_evaluation_id ON student_answers(evaluation_id);",
        "CREATE INDEX IF NOT EXISTS idx_student_answers_access_code ON student_answers(access_code);",
        "CREATE INDEX IF NOT EXISTS idx_student_answers_email_aluno ON student_answers(email_aluno);",
    ]

    with db_connect() as conn:
        with conn.cursor() as cur:
            cur.execute(ddl)
            cur.execute(student_answers_ddl)
            cur.execute("ALTER TABLE evaluations ADD COLUMN IF NOT EXISTS access_code TEXT;")
            cur.execute("ALTER TABLE evaluations ALTER COLUMN email_aluno SET DEFAULT '';")
            cur.execute("UPDATE evaluations SET access_code = upper(substr(replace(id, '-', ''), 1, 8)) WHERE access_code IS NULL OR access_code = '';")
            for idx in indexes:
                cur.execute(idx)
        conn.commit()


def db_insert(row: dict) -> None:
    cols = list(row.keys())
    placeholders = ", ".join(["%s"] * len(cols))
    columns_sql = ", ".join(cols)
    values = [row[c] for c in cols]

    sql = f"INSERT INTO evaluations ({columns_sql}) VALUES ({placeholders});"

    with db_connect() as conn:
        with conn.cursor() as cur:
            cur.execute(sql, values)
        conn.commit()


def db_list() -> pd.DataFrame:
    sql = "SELECT * FROM evaluations ORDER BY created_at DESC;"
    with db_connect() as conn:
        with conn.cursor() as cur:
            cur.execute(sql)
            rows = cur.fetchall()
    return pd.DataFrame(rows)


def db_get_by_id(evaluation_id: str) -> Optional[dict]:
    sql = "SELECT * FROM evaluations WHERE id = %s;"
    with db_connect() as conn:
        with conn.cursor() as cur:
            cur.execute(sql, (evaluation_id,))
            row = cur.fetchone()
    return dict(row) if row else None


def db_get_by_access_code(access_code: str) -> Optional[dict]:
    sql = "SELECT * FROM evaluations WHERE upper(access_code) = upper(%s);"
    with db_connect() as conn:
        with conn.cursor() as cur:
            cur.execute(sql, (access_code,))
            row = cur.fetchone()
    return dict(row) if row else None


def db_update_status(evaluation_id: str, new_status: str, observacao_admin: str) -> None:
    sql = """
    UPDATE evaluations
    SET status = %s,
        observacao_admin = %s,
        updated_at = %s
    WHERE id = %s;
    """
    with db_connect() as conn:
        with conn.cursor() as cur:
            cur.execute(sql, (new_status, observacao_admin, datetime.now(), evaluation_id))
        conn.commit()


def db_insert_student_answer(row: dict) -> None:
    cols = list(row.keys())
    placeholders = ", ".join(["%s"] * len(cols))
    columns_sql = ", ".join(cols)
    values = [row[c] for c in cols]

    sql = f"INSERT INTO student_answers ({columns_sql}) VALUES ({placeholders});"

    with db_connect() as conn:
        with conn.cursor() as cur:
            cur.execute(sql, values)
        conn.commit()


def db_list_student_answers() -> pd.DataFrame:
    sql = """
    SELECT
        sa.*,
        ev.titulo,
        ev.nome_docente,
        ev.email_docente,
        ev.area_conhecimento,
        ev.avaliacao,
        ev.bimestre,
        ev.turma
    FROM student_answers sa
    JOIN evaluations ev ON ev.id = sa.evaluation_id
    ORDER BY sa.created_at DESC;
    """
    with db_connect() as conn:
        with conn.cursor() as cur:
            cur.execute(sql)
            rows = cur.fetchall()
    return pd.DataFrame(rows)


def db_get_evaluation_for_answer(answer_id: str) -> Optional[dict]:
    sql = """
    SELECT ev.*
    FROM student_answers sa
    JOIN evaluations ev ON ev.id = sa.evaluation_id
    WHERE sa.id = %s;
    """
    with db_connect() as conn:
        with conn.cursor() as cur:
            cur.execute(sql, (answer_id,))
            row = cur.fetchone()
    return dict(row) if row else None


# ============================================================
# DOCX
# ============================================================

def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_centered_picture_if_exists(doc: Document, image_path: Path, width: float = 3.5) -> None:
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(width))


def add_question_image_if_exists(doc: Document, image_path_value: str, width: float = 5.5) -> None:
    if not str(image_path_value or "").strip():
        return
    image_path = Path(image_path_value or "")
    if image_path.exists() and image_path.is_file():
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(image_path), width=Inches(width))
        except Exception:
            doc.add_paragraph(f"[Figura não inserida no DOCX: arquivo de imagem inválido ou incompatível - {image_path.name}]")


def add_key_value_table(doc: Document, pairs: list[tuple[str, str]], cols: int = 2) -> None:
    rows = (len(pairs) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    for idx, (label, value) in enumerate(pairs):
        cell = table.cell(idx // cols, idx % cols)
        p = cell.paragraphs[0]
        p.add_run(f"{label}: ").bold = True
        p.add_run(value or "")


def add_response_lines(doc: Document, total: int) -> None:
    for i in range(1, total + 1):
        doc.add_paragraph(f"{i} " + "_" * 86)


def make_docx_bytes(fields: dict, created_at: str, sub_id: str, status: str) -> bytes:
    doc = Document()
    set_default_font(doc)

    add_centered_picture_if_exists(doc, BASE_DIR / "assets" / "cabecalho_evento.png")

    title = doc.add_heading("AVALIAÇÃO TEÓRICA-COGNITIVA", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_key_value_table(
        doc,
        [
            ("Área de conhecimento", fields.get("area_conhecimento", "")),
            ("Avaliação", fields.get("avaliacao", "")),
            ("Bimestre", fields.get("bimestre", "")),
            ("Turma", fields.get("turma", "")),
            ("Período letivo", fields.get("periodo_letivo", "")),
            ("Data", fields.get("data_avaliacao", "")),
            ("Nome do Docente", fields.get("nome_docente", "")),
            ("E-mail do Docente", fields.get("email_docente", "")),
            ("Código da avaliação", fields.get("access_code", "")),
            ("Nota da prova", ""),
        ],
    )

    doc.add_paragraph()
    doc.add_heading("ANTES DE INICIAR A PROVA LEIA ATENTAMENTE AS INSTRUÇÕES ABAIXO", level=2)
    for item in INSTRUCOES_PROVA:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "O valor total da prova é de "
        f"{fields.get('valor_total', '')}. "
        "A1 e A2: mínimo de 6,0 e máximo de 8,0 pontos. Na A3 o valor deve ser de 10 pontos."
    )

    for idx, questao in enumerate(fields["questoes"], start=1):
        doc.add_page_break()
        doc.add_heading(f"Questão {idx}    Valor: {questao.get('valor', '')}", level=2)
        doc.add_paragraph(
            f"Área do conhecimento: {questao.get('area', '')}    "
            f"Taxonomia: {questao.get('taxonomia', '')}"
        )
        doc.add_paragraph(questao.get("enunciado", ""))
        add_question_image_if_exists(doc, questao.get("image_path", ""))

        add_response_lines(doc, int(questao.get("linhas_resposta") or 8))

        doc.add_paragraph()
        doc.add_heading("Dimensões de avaliação qualitativa processual", level=3)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = "DIMENSÕES"
        table.cell(0, 1).text = "Registro"
        for dimensao in DIMENSOES_AVALIACAO:
            row = table.add_row().cells
            row[0].text = dimensao
            row[1].text = ""
        row = table.add_row().cells
        row[0].text = "Nota final atribuída"
        row[1].text = ""

        doc.add_heading("Gabarito", level=3)
        doc.add_paragraph(questao.get("gabarito", ""))
        doc.add_heading("Referência", level=3)
        doc.add_paragraph(questao.get("referencia", ""))

    doc.add_page_break()
    doc.add_heading("Metadados", level=2)
    doc.add_paragraph(
        f"Título interno: {fields.get('titulo', '')}\n"
        f"Data de cadastro: {created_at}\n"
        f"Protocolo: {sub_id}\n"
        f"Status: {status}"
    )

    buff = BytesIO()
    doc.save(buff)
    return buff.getvalue()


def make_student_answer_docx_bytes(evaluation: dict, student: dict, respostas: list[dict], created_at: str, answer_id: str) -> bytes:
    doc = Document()
    set_default_font(doc)

    add_centered_picture_if_exists(doc, BASE_DIR / "assets" / "cabecalho_evento.png")

    title = doc.add_heading("RESPOSTAS DA AVALIAÇÃO TEÓRICA-COGNITIVA", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_key_value_table(
        doc,
        [
            ("Aluno", student.get("nome_aluno", "")),
            ("E-mail do aluno", student.get("email_aluno", "")),
            ("Docente", evaluation.get("nome_docente", "")),
            ("Área de conhecimento", evaluation.get("area_conhecimento", "")),
            ("Avaliação", evaluation.get("avaliacao", "")),
            ("Bimestre", evaluation.get("bimestre", "")),
            ("Turma", evaluation.get("turma", "")),
            ("Código", evaluation.get("access_code", "")),
        ],
    )

    questoes = evaluation.get("questoes") or []
    for idx, resposta in enumerate(respostas, start=1):
        questao = questoes[idx - 1] if idx - 1 < len(questoes) else {}
        doc.add_page_break()
        doc.add_heading(f"Questão {idx}", level=2)
        doc.add_paragraph(f"Taxonomia: {questao.get('taxonomia', '')}    Valor: {questao.get('valor', '')}")
        doc.add_paragraph("Enunciado:")
        doc.add_paragraph(questao.get("enunciado", ""))
        add_question_image_if_exists(doc, questao.get("image_path", ""))
        doc.add_paragraph("Resposta do aluno:")
        doc.add_paragraph(resposta.get("resposta", ""))

    doc.add_page_break()
    doc.add_heading("Metadados", level=2)
    doc.add_paragraph(
        f"Data de envio: {created_at}\n"
        f"Protocolo da resposta: {answer_id}\n"
        f"Protocolo da avaliação: {evaluation.get('id', '')}"
    )

    buff = BytesIO()
    doc.save(buff)
    return buff.getvalue()


def make_xlsx_bytes(df_: pd.DataFrame) -> bytes:
    buff = BytesIO()
    with pd.ExcelWriter(buff, engine="openpyxl") as writer:
        df_.to_excel(writer, index=False, sheet_name="Avaliacoes")
    return buff.getvalue()


def make_zip_all(df_: pd.DataFrame) -> bytes:
    buff = BytesIO()
    with ZipFile(buff, "w", compression=ZIP_DEFLATED) as z:
        z.writestr("avaliacoes.csv", df_.to_csv(index=False).encode("utf-8"))
        z.writestr("avaliacoes.xlsx", make_xlsx_bytes(df_))

        for _, r in df_.iterrows():
            protocolo = r.get("id", "")
            titulo = safe_filename(r.get("titulo", "") or "", max_len=70)
            created_txt = str(r.get("created_at", ""))[:10]
            base = f"{created_txt}_{titulo}_{(protocolo or '')[:8]}".strip("_")
            docxp = r.get("docx_path", "")

            if docxp and Path(docxp).exists():
                z.writestr(f"DOCX/{base}.docx", read_file_bytes(docxp))
            else:
                z.writestr(f"DOCX/NAO_ENCONTRADO_{base}.txt", f"DOCX não encontrado: {docxp}".encode("utf-8"))

    return buff.getvalue()


# ============================================================
# EMAIL
# ============================================================

def smtp_is_configured() -> bool:
    return bool(SMTP_HOST and SMTP_PORT and SMTP_USER and SMTP_PASSWORD and SMTP_FROM)


def guess_mime(filename: str) -> tuple[str, str]:
    fn = (filename or "").lower()
    if fn.endswith(".docx"):
        return ("application", "vnd.openxmlformats-officedocument.wordprocessingml.document")
    if fn.endswith(".xlsx"):
        return ("application", "vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    return ("application", "octet-stream")


def send_email_with_attachments(
    *,
    to_email: str,
    cc_email: Optional[str],
    subject: str,
    body: str,
    attachments: list[tuple[str, bytes]],
) -> None:
    msg = EmailMessage()
    msg["From"] = SMTP_FROM
    msg["To"] = to_email
    if cc_email:
        msg["Cc"] = cc_email
    msg["Subject"] = subject
    msg.set_content(body)

    for filename, file_bytes in attachments:
        maintype, subtype = guess_mime(filename)
        msg.add_attachment(file_bytes, maintype=maintype, subtype=subtype, filename=filename)

    with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=30) as server:
        server.ehlo()
        server.starttls()
        server.ehlo()
        server.login(SMTP_USER, SMTP_PASSWORD)
        server.send_message(msg)


def send_status_update_email(row: dict, new_status: str, observacao_admin: str) -> None:
    if not smtp_is_configured():
        return

    subject = f"[Humanitas] Atualização da avaliação - Protocolo {row['id'][:8]}"
    body = (
        f"Olá!\n\n"
        f"Houve uma atualização na avaliação discursiva.\n\n"
        f"Protocolo: {row['id']}\n"
        f"Código de acesso: {row.get('access_code', '')}\n"
        f"Título: {row['titulo']}\n"
        f"Área de conhecimento: {row['area_conhecimento']}\n"
        f"Novo status: {new_status}\n"
        f"Última atualização: {now_str()}\n\n"
    )
    if observacao_admin.strip():
        body += f"Observação:\n{observacao_admin.strip()}\n\n"
    body += "Atenciosamente,\nFaculdade Humanitas\n"

    send_email_with_attachments(
        to_email=row["email_docente"],
        cc_email=None,
        subject=subject,
        body=body,
        attachments=[],
    )


# ============================================================
# VALIDATION
# ============================================================

def validate_all(fields: dict) -> Tuple[bool, Dict[str, str]]:
    errors: Dict[str, str] = {}

    required = {
        "nome_docente": "Informe o nome do docente.",
        "email_docente": "Informe o e-mail do docente.",
        "area_conhecimento": "Informe a área de conhecimento.",
        "avaliacao": "Selecione a avaliação.",
        "bimestre": "Selecione o bimestre.",
        "turma": "Informe a turma.",
        "periodo_letivo": "Informe o período letivo.",
        "data_avaliacao": "Informe a data da avaliação.",
        "valor_total": "Informe o valor total da prova.",
        "titulo": "Informe um título interno para identificação.",
    }

    for key, message in required.items():
        if not str(fields.get(key, "")).strip():
            errors[key] = message

    if fields.get("email_docente") and not is_email_valid(fields["email_docente"]):
        errors["email_docente"] = "Informe um e-mail de docente válido."

    total = 0
    for idx, questao in enumerate(fields.get("questoes", []), start=1):
        prefix = f"questao_{idx}"
        for key, label in [
            ("valor", "valor"),
            ("area", "área do conhecimento"),
            ("taxonomia", "taxonomia"),
            ("enunciado", "enunciado"),
            ("gabarito", "gabarito"),
            ("referencia", "referência"),
        ]:
            if not str(questao.get(key, "")).strip():
                errors[f"{prefix}_{key}"] = f"Questão {idx}: informe {label}."

        image_file = questao.get("image_file")
        if image_file is not None and not allowed_image_extension(image_file.name):
            errors[f"{prefix}_imagem"] = f"Questão {idx}: a imagem deve ser PNG, JPG ou JPEG."
        elif image_file is not None and not has_supported_image_signature(image_file.getbuffer().tobytes()):
            errors[f"{prefix}_imagem"] = f"Questão {idx}: o arquivo enviado não parece ser uma imagem PNG/JPG válida."

        total += (
            count_chars(questao.get("enunciado", ""))
            + count_chars(questao.get("gabarito", ""))
            + count_chars(questao.get("referencia", ""))
        )

    if len(fields.get("questoes", [])) != 6:
        errors["questoes"] = "A avaliação deve conter exatamente 6 questões."
    return (len(errors) == 0), errors


def validate_student_answer(student: dict, respostas: list[dict]) -> Tuple[bool, Dict[str, str]]:
    errors: Dict[str, str] = {}

    if not student.get("nome_aluno", "").strip():
        errors["nome_aluno"] = "Informe seu nome completo."
    if not student.get("email_aluno", "").strip() or not is_email_valid(student.get("email_aluno", "")):
        errors["email_aluno"] = "Informe um e-mail válido."

    for idx, resposta in enumerate(respostas, start=1):
        if not resposta.get("resposta", "").strip():
            errors[f"resposta_{idx}"] = f"Responda a questão {idx}."

    return (len(errors) == 0), errors


# ============================================================
# UI
# ============================================================

def area(label: str, key: str, height: int = 180, max_chars: int | None = None) -> str:
    txt = st.text_area(label, key=key, height=height)
    n = count_chars(txt)
    if max_chars is None:
        st.caption(f"Caracteres: {n}")
    elif n <= max_chars:
        st.caption(f"Caracteres: {n}/{max_chars}")
    else:
        st.error(f"Caracteres: {n}/{max_chars} (excedeu o limite)")
    return txt


def render_question_form(idx: int, default_taxonomia: str) -> dict:
    st.markdown(f"### Questão {idx}")
    c1, c2, c3, c4 = st.columns([1, 2, 1, 1])
    with c1:
        valor = st.text_input("Valor", key=f"q{idx}_valor", placeholder="Ex.: 1,5")
    with c2:
        area_conhecimento = st.text_input("Área do conhecimento", key=f"q{idx}_area")
    with c3:
        taxonomia = st.selectbox(
            "Taxonomia",
            TAXONOMIA_OPTIONS,
            index=TAXONOMIA_OPTIONS.index(default_taxonomia),
            key=f"q{idx}_taxonomia",
        )
    with c4:
        linhas_resposta = st.number_input(
            "Linhas",
            min_value=2,
            max_value=20,
            value=14 if idx == 1 else 8 if idx in [2, 3] else 2,
            step=1,
            key=f"q{idx}_linhas",
        )

    enunciado = area("Questão completa", f"q{idx}_enunciado", height=220)
    image_file = st.file_uploader(
        "Figura da questão (opcional)",
        type=["png", "jpg", "jpeg"],
        accept_multiple_files=False,
        key=f"q{idx}_image",
    )
    if image_file is not None:
        try:
            st.image(image_file, caption=f"Figura da questão {idx}", use_container_width=True)
        except Exception:
            st.warning("Não foi possível pré-visualizar esta imagem. Envie um PNG ou JPG/JPEG válido.")
    col_gab, col_ref = st.columns(2)
    with col_gab:
        gabarito = area("Gabarito esperado", f"q{idx}_gabarito", height=150)
    with col_ref:
        referencia = area("Referência em norma Vancouver", f"q{idx}_referencia", height=150)

    return {
        "valor": valor,
        "area": area_conhecimento,
        "taxonomia": taxonomia,
        "linhas_resposta": int(linhas_resposta),
        "enunciado": enunciado,
        "image_file": image_file,
        "image_filename": "",
        "image_path": "",
        "gabarito": gabarito,
        "referencia": referencia,
    }


def render_protocol_search():
    st.subheader("Consulta pública por protocolo")
    protocolo_busca = st.text_input("Digite o protocolo completo", key="busca_protocolo")

    if st.button("Consultar protocolo"):
        protocolo_busca = (protocolo_busca or "").strip()
        if not protocolo_busca:
            st.warning("Informe o protocolo.")
            return

        row = db_get_by_id(protocolo_busca)
        if not row:
            st.error("Protocolo não encontrado.")
            return

        st.success("Avaliação encontrada.")
        st.write(f"**Protocolo:** {row['id']}")
        st.write(f"**Título:** {row['titulo']}")
        st.write(f"**Docente:** {row['nome_docente']}")
        st.write(f"**Área de conhecimento:** {row.get('area_conhecimento', '')}")
        st.write(f"**Turma:** {row.get('turma', '')}")
        st.write(f"**Status:** {row['status']}")
        st.write(f"**Data de cadastro:** {row['created_at']}")
        st.write(f"**Última atualização:** {row.get('updated_at', row['created_at'])}")
        if (row.get("observacao_admin") or "").strip():
            st.write(f"**Observação:** {row['observacao_admin']}")


def render_student_page():
    st.subheader("Página do aluno")
    access_code = st.text_input("Código da avaliação", key="student_access_code").strip().upper()

    if not access_code:
        st.info("Digite o código informado pelo professor para abrir a avaliação.")
        return

    evaluation = db_get_by_access_code(access_code)
    if not evaluation:
        st.error("Código não encontrado.")
        return

    if evaluation.get("status") not in ["APROVADA", "ENVIADA AO ALUNO"]:
        st.warning("Esta avaliação ainda não está liberada para resposta.")
        return

    st.success("Avaliação carregada.")
    st.write(f"**Título:** {evaluation['titulo']}")
    st.write(f"**Docente:** {evaluation['nome_docente']}")
    st.write(f"**Área:** {evaluation['area_conhecimento']}")
    st.write(f"**Turma:** {evaluation['turma']}")
    st.write(f"**Avaliação:** {evaluation['avaliacao']} | **Bimestre:** {evaluation['bimestre']}")
    st.write(f"**Data:** {evaluation['data_avaliacao']} | **Valor total:** {evaluation['valor_total']}")

    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        nome_aluno = st.text_input("Nome completo do aluno", key="student_nome_aluno")
    with col2:
        email_aluno = st.text_input("E-mail do aluno", key="student_email_aluno")

    st.divider()
    questoes = evaluation.get("questoes") or []
    respostas = []
    for idx, questao in enumerate(questoes, start=1):
        st.markdown(f"### Questão {idx}")
        st.caption(f"Taxonomia: {questao.get('taxonomia', '')} | Valor: {questao.get('valor', '')}")
        st.write(questao.get("enunciado", ""))
        if questao.get("image_path") and Path(questao["image_path"]).exists():
            try:
                st.image(questao["image_path"], caption=f"Figura da questão {idx}", use_container_width=True)
            except Exception:
                st.warning(f"A figura da questão {idx} não pôde ser exibida.")
        resposta = st.text_area("Resposta", key=f"student_resposta_{idx}", height=220)
        respostas.append({"questao": idx, "resposta": resposta})

    aceite = st.checkbox("Declaro que revisei minhas respostas e confirmo o envio.", key="student_aceite")

    if st.button("Enviar respostas"):
        student = {
            "nome_aluno": nome_aluno or "",
            "email_aluno": email_aluno or "",
        }
        ok, errors = validate_student_answer(student, respostas)
        if not aceite:
            ok = False
            errors["aceite"] = "Você precisa confirmar o envio."

        if not ok:
            st.error("Não foi possível enviar. Corrija os itens abaixo:")
            for k, v in errors.items():
                st.write(f"- **{k}**: {v}")
            st.stop()

        answer_id = str(uuid.uuid4())
        now = datetime.now()
        now_fmt = now.strftime("%Y-%m-%d %H:%M:%S")
        docx_name = f"{now_fmt[:10]}_{safe_slug(student['email_aluno'])}_{safe_slug(evaluation['titulo'])}_{answer_id[:8]}.docx"
        docx_path = DOCX_DIR / docx_name
        docx_bytes = make_student_answer_docx_bytes(evaluation, student, respostas, now_fmt, answer_id)
        docx_path.write_bytes(docx_bytes)

        row = {
            "id": answer_id,
            "evaluation_id": evaluation["id"],
            "access_code": evaluation["access_code"],
            "created_at": now,
            "updated_at": now,
            "nome_aluno": student["nome_aluno"].strip(),
            "email_aluno": student["email_aluno"].strip(),
            "respostas": Json(respostas),
            "docx_filename": docx_name,
            "docx_path": str(docx_path),
            "status": "ENVIADA",
            "observacao_admin": "",
        }
        db_insert_student_answer(row)

        st.success("Respostas enviadas com sucesso.")
        st.info(f"Protocolo da resposta: **{answer_id}**")
        st.download_button(
            "Baixar comprovante DOCX",
            data=docx_bytes,
            file_name=safe_filename(docx_name),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        if smtp_is_configured():
            subject = f"[Humanitas] Respostas recebidas - {evaluation['titulo']}"
            body = (
                f"Olá!\n\n"
                f"As respostas do aluno foram recebidas.\n\n"
                f"Aluno: {student['nome_aluno'].strip()}\n"
                f"E-mail: {student['email_aluno'].strip()}\n"
                f"Avaliação: {evaluation['titulo']}\n"
                f"Código: {evaluation['access_code']}\n"
                f"Protocolo da resposta: {answer_id}\n\n"
                f"Atenciosamente,\nFaculdade Humanitas\n"
            )
            try:
                send_email_with_attachments(
                    to_email=evaluation["email_docente"],
                    cc_email=student["email_aluno"],
                    subject=subject,
                    body=body,
                    attachments=[(docx_name, docx_bytes)],
                )
                st.success("Confirmação enviada por e-mail ao docente, com cópia para o aluno.")
            except Exception as e:
                st.warning(f"Resposta salva, mas houve falha no envio do e-mail: {e}")


def render_answer_comparison(answer_row: dict) -> None:
    evaluation = db_get_evaluation_for_answer(answer_row["id"])
    if not evaluation:
        st.warning("Não foi possível localizar a avaliação desta resposta.")
        return

    questoes = evaluation.get("questoes") or []
    respostas = answer_row.get("respostas") or []

    st.markdown("#### Correção comparativa")
    st.write(f"**Aluno:** {answer_row['nome_aluno']}")
    st.write(f"**Avaliação:** {evaluation['titulo']}")

    for idx, questao in enumerate(questoes, start=1):
        resposta = ""
        if idx - 1 < len(respostas):
            resposta = respostas[idx - 1].get("resposta", "")

        with st.expander(f"Questão {idx} | Taxonomia: {questao.get('taxonomia', '')} | Valor: {questao.get('valor', '')}", expanded=idx == 1):
            st.markdown("**Enunciado**")
            st.write(questao.get("enunciado", ""))
            if questao.get("image_path") and Path(questao["image_path"]).exists():
                try:
                    st.image(questao["image_path"], caption=f"Figura da questão {idx}", use_container_width=True)
                except Exception:
                    st.warning("A figura desta questão não pôde ser exibida.")

            col_resp, col_gab = st.columns(2)
            with col_resp:
                st.markdown("**Resposta do aluno**")
                st.write(resposta or "Sem resposta registrada.")
            with col_gab:
                st.markdown("**Gabarito esperado**")
                st.write(questao.get("gabarito", "") or "Sem gabarito registrado.")

            if questao.get("referencia"):
                st.markdown("**Referência**")
                st.write(questao.get("referencia", ""))


def main():
    st.set_page_config(page_title=APP_TITLE, page_icon="📝", layout="wide")
    db_init()

    logo_path = BASE_DIR / "assets" / "logo_humanitas.png"
    if logo_path.exists():
        c1, c2, c3 = st.columns([1, 2, 1])
        with c2:
            st.image(str(logo_path), width=320)

    st.markdown("<h1 style='text-align: center;'>Avaliação Teórica-Cognitiva</h1>", unsafe_allow_html=True)
    st.markdown(
        "<p style='text-align: center;'>Aplicativo para avaliação discursiva com 6 questões</p>",
        unsafe_allow_html=True,
    )

    tab1, tab2, tab3, tab4 = st.tabs(["Professor", "Aluno", "Consulta por protocolo", "Admin"])

    with tab1:
        st.subheader("Dados da avaliação")
        col1, col2 = st.columns(2)
        with col1:
            nome_docente = st.text_input("Nome do docente", key="nome_docente")
            email_docente = st.text_input("E-mail do docente", key="email_docente")
            area_conhecimento = st.text_input("Área de conhecimento", key="area_conhecimento")
        with col2:
            avaliacao = st.selectbox("Avaliação", AVALIACAO_OPTIONS, key="avaliacao")
            bimestre = st.selectbox("Bimestre", BIMESTRE_OPTIONS, key="bimestre")
            turma = st.text_input("Turma", key="turma", placeholder="Ex.: MED 3A")
            periodo_letivo = st.text_input("Período letivo", key="periodo_letivo", value="2026-1")
            data_avaliacao = st.text_input("Data da avaliação", key="data_avaliacao", placeholder="____/____/2026")
            valor_total = st.text_input("Valor total da prova", key="valor_total", placeholder="Ex.: 8,0")

        titulo = st.text_input(
            "Título interno para identificação",
            key="titulo",
            placeholder="Ex.: A1 - Bases Morfofuncionais - Turma MED 3A",
        )

        st.divider()
        st.subheader("Questões discursivas")
        st.caption("Sequência sugerida pelo modelo: Questão 1 Alta; Questões 2 e 3 Média; Questões 4, 5 e 6 Baixa.")

        defaults = ["Alta", "Média", "Média", "Baixa", "Baixa", "Baixa"]
        questoes = []
        for idx, default_taxonomia in enumerate(defaults, start=1):
            with st.expander(f"Questão {idx} - Taxonomia {default_taxonomia}", expanded=idx == 1):
                questoes.append(render_question_form(idx, default_taxonomia))

        total_chars = sum(
            count_chars(q["enunciado"]) + count_chars(q["gabarito"]) + count_chars(q["referencia"])
            for q in questoes
        )
        st.info(f"Conteúdo total: {total_chars} caracteres")

        st.divider()
        aceite = st.checkbox("Declaro que revisei a avaliação e confirmo o cadastro.")

        if st.button("Gerar avaliação"):
            fields = {
                "nome_docente": nome_docente or "",
                "email_docente": email_docente or "",
                "area_conhecimento": area_conhecimento or "",
                "avaliacao": avaliacao or "",
                "bimestre": bimestre or "",
                "turma": turma or "",
                "periodo_letivo": periodo_letivo or "",
                "data_avaliacao": data_avaliacao or "",
                "valor_total": valor_total or "",
                "titulo": titulo or "",
                "questoes": questoes,
            }

            ok, errors = validate_all(fields)
            if not aceite:
                ok = False
                errors["aceite"] = "Você precisa marcar o aceite para concluir o cadastro."

            if not ok:
                st.error("Não foi possível cadastrar. Corrija os itens abaixo:")
                for k, v in errors.items():
                    st.write(f"- **{k}**: {v}")
                st.stop()

            sub_id = str(uuid.uuid4())
            access_code = generate_access_code()
            now = datetime.now()
            now_fmt = now.strftime("%Y-%m-%d %H:%M:%S")
            status = "RECEBIDA"
            fields["access_code"] = access_code

            questoes_para_salvar = []
            for idx, questao in enumerate(fields["questoes"], start=1):
                questao_limpa = {k: v for k, v in questao.items() if k != "image_file"}
                image_file = questao.get("image_file")
                if image_file is not None:
                    image_name, image_path = save_question_image(image_file, sub_id, idx)
                    questao_limpa["image_filename"] = image_name
                    questao_limpa["image_path"] = image_path
                questoes_para_salvar.append(questao_limpa)
            fields["questoes"] = questoes_para_salvar

            docx_name = f"{now_fmt[:10]}_{safe_slug(fields['email_docente'])}_{safe_slug(fields['titulo'])}_{sub_id[:8]}.docx"
            docx_path = DOCX_DIR / docx_name
            docx_bytes = make_docx_bytes(fields, created_at=now_fmt, sub_id=sub_id, status=status)
            docx_path.write_bytes(docx_bytes)

            row = {
                "id": sub_id,
                "created_at": now,
                "updated_at": now,
                "nome_docente": fields["nome_docente"].strip(),
                "email_docente": fields["email_docente"].strip(),
                "email_aluno": "",
                "area_conhecimento": fields["area_conhecimento"].strip(),
                "avaliacao": fields["avaliacao"].strip(),
                "bimestre": fields["bimestre"].strip(),
                "turma": fields["turma"].strip(),
                "periodo_letivo": fields["periodo_letivo"].strip(),
                "data_avaliacao": fields["data_avaliacao"].strip(),
                "valor_total": fields["valor_total"].strip(),
                "titulo": fields["titulo"].strip(),
                "access_code": access_code,
                "questoes": Json(fields["questoes"]),
                "docx_filename": docx_name,
                "docx_path": str(docx_path),
                "status": status,
                "observacao_admin": "",
            }
            db_insert(row)

            st.success("Avaliação cadastrada com sucesso.")
            st.info(f"Protocolo: **{sub_id}**")
            st.info(f"Código para disponibilizar aos alunos após aprovação: **{access_code}**")

            st.download_button(
                "Baixar DOCX gerado agora",
                data=docx_bytes,
                file_name=safe_filename(docx_name),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            if smtp_is_configured():
                subject = f"[Humanitas] Avaliação teórica-cognitiva - Protocolo {sub_id[:8]}"
                body = (
                    f"Olá!\n\n"
                    f"Segue em anexo a avaliação teórica-cognitiva cadastrada.\n\n"
                    f"Protocolo: {sub_id}\n"
                    f"Código de acesso do aluno: {access_code}\n"
                    f"Docente: {fields['nome_docente'].strip()}\n"
                    f"Área de conhecimento: {fields['area_conhecimento'].strip()}\n"
                    f"Avaliação: {fields['avaliacao'].strip()}\n"
                    f"Bimestre: {fields['bimestre'].strip()}\n"
                    f"Turma: {fields['turma'].strip()}\n"
                    f"Data: {fields['data_avaliacao'].strip()}\n\n"
                    f"Atenciosamente,\nFaculdade Humanitas\n"
                )
                try:
                    send_email_with_attachments(
                        to_email=fields["email_docente"].strip(),
                        cc_email=None,
                        subject=subject,
                        body=body,
                        attachments=[(docx_name, docx_bytes)],
                    )
                    st.success("E-mail enviado ao docente.")
                except Exception as e:
                    st.error(f"Falha ao enviar e-mail automático: {e}")

    with tab2:
        render_student_page()

    with tab3:
        render_protocol_search()

    with tab4:
        st.subheader("Painel administrativo")

        if not admin_auth_guard(ADMIN_PASSWORD):
            st.stop()

        if st.button("Sair do Admin"):
            st.session_state.admin_authed = False
            st.rerun()

        df = db_list()

        if len(df) == 0:
            st.info("Nenhuma avaliação cadastrada ainda.")
            st.stop()

        st.write(f"Total de avaliações: **{len(df)}**")

        f1, f2, f3 = st.columns(3)
        with f1:
            filtro_status = st.selectbox("Filtrar por status", ["Todos"] + STATUS_OPTIONS)
        with f2:
            turmas = sorted([x for x in df["turma"].fillna("").unique().tolist() if x])
            filtro_turma = st.selectbox("Filtrar por turma", ["Todos"] + turmas)
        with f3:
            areas = sorted([x for x in df["area_conhecimento"].fillna("").unique().tolist() if x])
            filtro_area = st.selectbox("Filtrar por área", ["Todos"] + areas)

        busca_texto = st.text_input("Buscar por docente, título, código ou protocolo")
        df_f = df.copy()

        if filtro_status != "Todos":
            df_f = df_f[df_f["status"] == filtro_status]
        if filtro_turma != "Todos":
            df_f = df_f[df_f["turma"] == filtro_turma]
        if filtro_area != "Todos":
            df_f = df_f[df_f["area_conhecimento"] == filtro_area]
        if (busca_texto or "").strip():
            termo = busca_texto.strip().lower()
            mask = (
                df_f["nome_docente"].fillna("").astype(str).str.lower().str.contains(termo, regex=False)
                | df_f["email_docente"].fillna("").astype(str).str.lower().str.contains(termo, regex=False)
                | df_f["titulo"].fillna("").astype(str).str.lower().str.contains(termo, regex=False)
                | df_f["access_code"].fillna("").astype(str).str.lower().str.contains(termo, regex=False)
                | df_f["id"].fillna("").astype(str).str.lower().str.contains(termo, regex=False)
            )
            df_f = df_f[mask]

        show_cols = [
            "created_at",
            "updated_at",
            "nome_docente",
            "email_docente",
            "area_conhecimento",
            "avaliacao",
            "bimestre",
            "turma",
            "titulo",
            "access_code",
            "status",
            "docx_filename",
            "id",
        ]
        st.dataframe(df_f[show_cols], use_container_width=True, hide_index=True)

        st.markdown("### Exportações")
        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button(
                "Baixar CSV",
                data=df_f.to_csv(index=False).encode("utf-8"),
                file_name="avaliacoes_filtradas.csv",
                mime="text/csv",
            )
        with c2:
            st.download_button(
                "Baixar Excel",
                data=make_xlsx_bytes(df_f),
                file_name="avaliacoes_filtradas.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        with c3:
            st.download_button(
                "Exportar ZIP completo",
                data=make_zip_all(df_f),
                file_name="exportacao_avaliacoes.zip",
                mime="application/zip",
            )

        st.markdown("### Respostas dos alunos")
        answers_df = db_list_student_answers()
        if len(answers_df) == 0:
            st.info("Nenhuma resposta de aluno enviada ainda.")
        else:
            answer_cols = [
                "created_at",
                "nome_aluno",
                "email_aluno",
                "titulo",
                "access_code",
                "turma",
                "status",
                "docx_filename",
                "id",
            ]
            st.dataframe(answers_df[answer_cols], use_container_width=True, hide_index=True)

            answer_options = []
            answer_map = {}
            for _, answer in answers_df.iterrows():
                label = (
                    f"{str(answer['created_at'])[:19]} | {answer['nome_aluno']} | "
                    f"{str(answer['titulo'])[:50]} | {str(answer['id'])[:8]}"
                )
                answer_options.append(label)
                answer_map[label] = answer["id"]

            selected_answer_label = st.selectbox("Selecionar resposta de aluno", answer_options)
            selected_answer_id = answer_map[selected_answer_label]
            answer_row = answers_df[answers_df["id"] == selected_answer_id].iloc[0].to_dict()

            st.write(f"**Aluno:** {answer_row['nome_aluno']}")
            st.write(f"**E-mail:** {answer_row['email_aluno']}")
            st.write(f"**Avaliação:** {answer_row['titulo']}")
            st.write(f"**Código:** {answer_row['access_code']}")

            if answer_row.get("docx_path") and Path(answer_row["docx_path"]).exists():
                st.download_button(
                    "Baixar resposta DOCX",
                    data=read_file_bytes(answer_row["docx_path"]),
                    file_name=safe_filename(answer_row["docx_filename"]),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_answer_docx_{answer_row['id']}",
                )
            else:
                st.warning("DOCX da resposta não encontrado.")

            render_answer_comparison(answer_row)

        st.markdown("### Gestão de avaliação")
        opcoes = []
        id_map = {}
        for _, r in df_f.iterrows():
            created_txt = str(r["created_at"])[:19]
            label = f"{created_txt} | {r['nome_docente']} | {str(r['titulo'])[:60]} | {str(r['id'])[:8]}"
            opcoes.append(label)
            id_map[label] = r["id"]

        if not opcoes:
            st.info("Nenhuma avaliação disponível com os filtros atuais.")
            st.stop()

        selected_label = st.selectbox("Selecione uma avaliação", opcoes)
        selected_id = id_map[selected_label]
        row = df[df["id"] == selected_id].iloc[0].to_dict()

        st.write(f"**Protocolo:** {row['id']}")
        st.write(f"**Código para alunos:** {row.get('access_code', '')}")
        st.write(f"**Docente:** {row['nome_docente']}")
        st.write(f"**E-mail do docente:** {row['email_docente']}")
        st.write(f"**Área:** {row.get('area_conhecimento', '')}")
        st.write(f"**Avaliação:** {row.get('avaliacao', '')}")
        st.write(f"**Turma:** {row.get('turma', '')}")
        st.write(f"**Título:** {row['titulo']}")
        st.write(f"**Status atual:** {row['status']}")

        if (row.get("observacao_admin") or "").strip():
            st.write(f"**Observação atual:** {row['observacao_admin']}")

        if row.get("docx_path") and Path(row["docx_path"]).exists():
            st.download_button(
                "Baixar DOCX",
                data=read_file_bytes(row["docx_path"]),
                file_name=safe_filename(row["docx_filename"]),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_docx_{row['id']}",
            )
        else:
            st.warning("DOCX local não encontrado.")

        st.markdown("#### Atualizar status")
        novo_status = st.selectbox(
            "Novo status",
            STATUS_OPTIONS,
            index=STATUS_OPTIONS.index(row["status"]) if row["status"] in STATUS_OPTIONS else 0,
        )
        observacao_admin = st.text_area(
            "Observação para o aluno/docente",
            value=row.get("observacao_admin", ""),
            height=120,
        )
        enviar_email_status = st.checkbox(
            "Enviar e-mail automático ao docente sobre a mudança de status",
            value=True,
        )

        if st.button("Salvar atualização de status"):
            mudou = (novo_status != row["status"]) or ((observacao_admin or "") != (row.get("observacao_admin", "") or ""))
            db_update_status(selected_id, novo_status, observacao_admin)

            if mudou and enviar_email_status:
                try:
                    row["status"] = novo_status
                    row["observacao_admin"] = observacao_admin
                    send_status_update_email(row, novo_status, observacao_admin)
                    st.success("Status atualizado e e-mail enviado ao docente.")
                except Exception as e:
                    st.warning(f"Status atualizado, mas houve falha no envio do e-mail: {e}")
            else:
                st.success("Status atualizado com sucesso.")

            st.rerun()


if __name__ == "__main__":
    main()
